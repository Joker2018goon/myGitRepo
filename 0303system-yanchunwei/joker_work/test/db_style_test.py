import os
import sys
dir_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(dir_path)
from docx.enum.style import WD_STYLE_TYPE
from docx import *
from utils.path_manage import PathManage

document = Document()
styles = document.styles

#生成所有表样式
for s in styles:
    if s.type == WD_STYLE_TYPE.TABLE:
        document.add_paragraph("Table style is :  " + s.name)
        document.add_table(3, 3, style=s)
        document.add_paragraph("\n")

document.save(PathManage.doc_path('db.docx'))
