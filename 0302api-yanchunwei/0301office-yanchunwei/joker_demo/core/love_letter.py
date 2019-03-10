# -*- coding:utf-8 -*-
# 根据de8ug老师讲的表白信中也是对word文档的操作，但是比较特殊，所有运用继承的思想来做题

import os
import sys
dir_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(dir_path)
import time
import logging
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
from core.base.word_base import Word
from utils.path_manage import PathManage
from utils.log_util import Logger

mylog=Logger(__name__).getlog()
# 关闭日志功能
# mylog.disable(logging.DEBUG)


class LoveLetter(Word):
    '''继承Word基类，个性化loveletter'''

    def __init__(self):
        '''初始化'''
        Word.__init__(self)

    def add_heart_method(self, sentence='love', style='Caption'):
        '''打印心形文字，借鉴网上代码'''
        try:
            # sentence = "love"
            for char in sentence.split():
                allChar = []
                for y in range(10, -10, -1):
                    lst = []
                    lst_con = ''
                    for x in range(-30, 35):
                        formula = (
                            (x * 0.04)**2 +
                            (y * 0.2)**2 - 1)**3 - (x * 0.04)**2 * (y * 0.2)**3
                        if formula <= 0:
                            lst_con += char[(x) % len(char)]
                        else:
                            lst_con += ' '
                    lst.append(lst_con)
                    allChar += lst
                # print('\n'.join(allChar))
                mylog.info('打印心形')
                mylog.info('\n'.join(allChar))
                p = self.add_paragraph_method('\n'.join(allChar), style)
                self.add_text_style(p)
                time.sleep(1)
        except Exception as e:
            # print(e)
            mylog.error('add_heart error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_heart error'

    def add_lover_method(self, lover='joker'):
        '''添加副标题以及索要表达的对象lover'''
        try:
            self.document.add_paragraph(f'Dear {lover}', 'Subtitle')
        except Exception as e:
            # print(e)
            mylog.error('add_lover error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_lover error'

    def add_sign_method(self, sign_name='joker'):
        '''添加署名'''
        try:
            sign = self.document.add_paragraph(f'Your {sign_name}', 'Subtitle')
            # 段落对齐,0=左对齐,1=居中,2=右对齐
            sign.alignment = 2
        except Exception as e:
            # print(e)
            mylog.error('add_sign error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_sign error'

    def add_title_method(self, title_context='LoveLetter',level=0):
        '''添加表发信大标题'''
        try:
            title = self.document.add_heading(title_context, level)
            # title.alignment=WD_UNDERLINE
            self.add_text_style(title)
        except Exception as e:
            # print(e)
            mylog.error('add_title error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_title error'


def main():
    love = LoveLetter()
    love.add_title_method('这是一个love模板')
    love.add_lover_method('玲玲')
    love.add_paragraph_method('正文' * 10)
    love.add_heart_method()
    love.add_sign_method('张三')
    love.save_method(PathManage.doc_path('love.docx'))


if __name__ == '__main__':
    main()
