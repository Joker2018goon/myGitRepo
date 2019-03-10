# !/usr/bin/env Python
# -*- coding:utf-8 -*-
# word_base.py word 基类

import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from utils.log_util import Logger

mylog=Logger(__name__).getlog()
# 关闭日志功能
# mylog.disable(logging.DEBUG)

class Word:
    '''word 基类，包含一些word文档的基本操作'''

    def __init__(self):
        '''初始化'''
        self.document = Document()
        self.ret = {'state': 0, 'stateMessage': 'success'}

    def save_method(self, file_name):
        '''保存文件'''
        try:
            self.document.save(file_name)
        except Exception:
            # print(e)
            mylog.error('save error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'save error'

    def add_head_method(self, str, level):
        '''添加不同等级的标题'''
        try:
            title=self.document.add_heading(str, level)
        except Exception:
            # print(e)
            mylog.error('add_head error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_head error'
        return title

    def add_paragraph_method(self, str, style=None):
        '''添加段落文本,可以带样式'''
        try:
            paragraph = self.document.add_paragraph(str, style)
        except Exception:
            # print(e)
            mylog.error('add_paragraph error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_paragraph error'
        return paragraph

    def add_paragraph_words_method(self, paragraph, str):
        '''对已有的段落，追加文本'''
        try:
            run_words = paragraph.add_run(str)
        except Exception:
            # print(e)
            mylog.error('add_paragraph_words error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_paragraph_words error'
        return run_words

    def add_style_method(self, text_str, style):
        '''对文本，追加样式'''
        try:
            # if style.type==WD_STYLE_TYPE:
            text_str.style = style
        except Exception:
            # print(e)
            mylog.error('add_style error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_style error'

    def add_fontName_method(self, text_str, style):
        '''对文本，设置字体'''
        try:
            # if style.type==WD_STYLE_TYPE:
            text_str.font.name = style
        except Exception:
            # print(e)
            mylog.error('add_fontName error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_fontName error'

    def add_fontSize_method(self, text_str, num):
        '''对文本，设置字号'''
        try:
            text_str.font.size = Pt(num)
        except Exception:
            # print(e)
            mylog.error('add_fontSize error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_fontSize error'

    def add_bold_method(self, text_str):
        '''加粗文本'''
        try:
            text_str.bold = True
        except Exception:
            # print(e)
            mylog.error('add_bold error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_bold error'

    def add_table_method(self, row, col):
        '''添加表格'''
        try:
            self.table = self.document.add_table(rows=row, cols=col)
        except Exception:
            # print(e)
            mylog.error('add_table error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_table error'

    def add_page_method(self):
        '''添加分页'''
        pass

    def add_pic_method(self, path_str):
        '''添加图片'''
        try:
            self.document.add_picture(path_str)
        except Exception:
            # print(e)
            mylog.error('add_pic error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_pic error'

    def add_text_style(self, paragraph):
        '''添加文本样式'''
        try:
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # 段落对齐,0=左对齐,1=居中,2=右对齐
            paragraph.alignment = 1
        except Exception:
            # print(e)
            mylog.error('add_pic error')
            self.ret['state'] = 1
            self.ret['stateMessage'] = 'add_pic error'