import os
import sys
dir_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(dir_path)

from docx.enum.style import WD_STYLE_TYPE
from docx import *
from utils.path_manage import PathManage

document = Document()
styles = document.styles
para = document.add_paragraph()

#生成所有字符样式
for s in styles:
    if s.type == WD_STYLE_TYPE.CHARACTER:
        run = para.add_run("Character style is:  "+s.name+"\n")
        run.style = s

document.save(PathManage.doc_path('character_style.docx'))
